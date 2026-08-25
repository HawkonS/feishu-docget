import os
import tempfile
import unittest

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from PIL import Image

from src.converters.docx.cleaner import clean_document


class TableParagraphStyleTests(unittest.TestCase):
    def _save_table_document(self, path, template_path=None):
        document = Document(template_path) if template_path else Document()
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = 'header'
        table.cell(0, 1).text = 'header'
        table.cell(1, 0).text = 'body'
        table.cell(1, 1).text = 'body'
        document.save(path)

    def test_table_paragraphs_use_separate_presets(self):
        with tempfile.TemporaryDirectory() as workspace:
            output_path = os.path.join(workspace, 'output.docx')
            self._save_table_document(output_path)
            clean_document(output_path)

            document = Document(output_path)
            styles = {style.name: style for style in document.styles}
            self.assertEqual(styles['表格正文'].type, WD_STYLE_TYPE.PARAGRAPH)
            self.assertEqual(styles['表格表头'].type, WD_STYLE_TYPE.PARAGRAPH)
            self.assertIsNot(styles['表格正文'], styles['表格表头'])
            self.assertEqual(document.tables[0].cell(0, 0).paragraphs[0].style.name, '表格表头')
            self.assertEqual(document.tables[0].cell(1, 0).paragraphs[0].style.name, '表格正文')

    def test_existing_paragraph_presets_are_reused(self):
        with tempfile.TemporaryDirectory() as workspace:
            template_path = os.path.join(workspace, 'template.docx')
            output_path = os.path.join(workspace, 'output.docx')
            template = Document()
            body_style = template.styles.add_style('表格正文', WD_STYLE_TYPE.PARAGRAPH)
            header_style = template.styles.add_style('表格表头', WD_STYLE_TYPE.PARAGRAPH)
            body_style.font.name = 'Courier New'
            header_style.font.name = 'Arial'
            template.save(template_path)
            self._save_table_document(output_path, template_path=template_path)

            clean_document(output_path, template_path=template_path)

            document = Document(output_path)
            names = [style.name for style in document.styles]
            self.assertEqual(names.count('表格正文'), 1)
            self.assertEqual(names.count('表格表头'), 1)
            self.assertNotIn('表格正文（导出）', names)
            self.assertNotIn('表格表头（导出）', names)

    def test_non_paragraph_name_collision_gets_reusable_suffix(self):
        with tempfile.TemporaryDirectory() as workspace:
            template_path = os.path.join(workspace, 'template.docx')
            output_path = os.path.join(workspace, 'output.docx')
            template = Document()
            template.styles.add_style('表格正文', WD_STYLE_TYPE.TABLE)
            template.styles.add_style('表格表头', WD_STYLE_TYPE.TABLE)
            template.save(template_path)
            self._save_table_document(output_path, template_path=template_path)

            clean_document(output_path, template_path=template_path)

            document = Document(output_path)
            names = [style.name for style in document.styles]
            self.assertEqual(names.count('表格正文'), 1)
            self.assertEqual(names.count('表格表头'), 1)
            self.assertEqual(names.count('表格正文（导出）'), 1)
            self.assertEqual(names.count('表格表头（导出）'), 1)
            self.assertEqual(
                document.tables[0].cell(1, 0).paragraphs[0].style.name,
                '表格正文（导出）',
            )
            self.assertEqual(
                document.tables[0].cell(0, 0).paragraphs[0].style.name,
                '表格表头（导出）',
            )

    def test_image_paragraph_uses_independent_preset(self):
        with tempfile.TemporaryDirectory() as workspace:
            image_path = os.path.join(workspace, 'image.png')
            output_path = os.path.join(workspace, 'output.docx')
            Image.new('RGB', (12, 12), color='red').save(image_path)

            document = Document()
            image_paragraph = document.add_paragraph()
            image_paragraph.add_run().add_picture(image_path)
            document.add_paragraph('caption')
            document.save(output_path)

            clean_document(
                output_path,
                body_style={
                    'fontSize': 30,
                    'lineSpacing': 2,
                    'spaceBefore': 4,
                    'spaceBeforeUnit': 'lines',
                },
            )

            document = Document(output_path)
            self.assertEqual(document.paragraphs[0].style.name, '图片')
            self.assertEqual(document.paragraphs[1].style.name, 'Normal')
            self.assertEqual(
                document.paragraphs[0]._element.pPr.find(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}spacing'
                ),
                None,
            )
            self.assertEqual(
                document.styles['图片'].type,
                WD_STYLE_TYPE.PARAGRAPH,
            )

    def test_existing_image_paragraph_preset_is_reused(self):
        with tempfile.TemporaryDirectory() as workspace:
            image_path = os.path.join(workspace, 'image.png')
            template_path = os.path.join(workspace, 'template.docx')
            output_path = os.path.join(workspace, 'output.docx')
            Image.new('RGB', (12, 12), color='blue').save(image_path)

            template = Document()
            template.styles.add_style('图片', WD_STYLE_TYPE.PARAGRAPH)
            template.save(template_path)
            document = Document(template_path)
            document.add_paragraph().add_run().add_picture(image_path)
            document.save(output_path)

            clean_document(output_path, template_path=template_path)

            document = Document(output_path)
            names = [style.name for style in document.styles]
            self.assertEqual(names.count('图片'), 1)
            self.assertNotIn('图片（导出）', names)
            self.assertEqual(document.paragraphs[0].style.name, '图片')


if __name__ == '__main__':
    unittest.main()
